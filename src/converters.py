import csv
import json
import os

from PIL import Image
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import openpyxl
from docx import Document

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".ico", ".tiff"}


def is_image(path):
    return os.path.splitext(path)[1].lower() in IMAGE_EXT


def convert_image(src, dst):
    img = Image.open(src)
    if dst.lower().endswith((".jpg", ".jpeg")) and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(dst)


def images_to_pdf(paths, dst):
    c = canvas.Canvas(dst, pagesize=A4)
    page_w, page_h = A4
    for p in paths:
        img = Image.open(p)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        w, h = img.size
        ratio = min(page_w / w, page_h / h)
        new_w, new_h = w * ratio, h * ratio
        x = (page_w - new_w) / 2
        y = (page_h - new_h) / 2
        c.drawImage(ImageReader(img), x, y, new_w, new_h)
        c.showPage()
    c.save()


def pdf_to_images(src, out_dir, fmt="png"):
    import pymupdf

    doc = pymupdf.open(src)
    saved = []
    base = os.path.splitext(os.path.basename(src))[0]
    for i, page in enumerate(doc, start=1):
        pix = page.get_pixmap(dpi=200)
        out_path = os.path.join(out_dir, f"{base}_page{i}.{fmt}")
        pix.save(out_path)
        saved.append(out_path)
    doc.close()
    return saved


def pdf_to_text(src, dst):
    reader = PdfReader(src)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
        text += "\n"
    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)


def text_to_pdf(src, dst):
    with open(src, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.read().splitlines()

    c = canvas.Canvas(dst, pagesize=A4)
    width, height = A4
    y = height - 50
    for line in lines:
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:110])
        y -= 14
    c.save()


def docx_to_text(src, dst):
    doc = Document(src)
    with open(dst, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")


def docx_to_pdf(src, dst):
    doc = Document(src)
    lines = [p.text for p in doc.paragraphs]
    tmp_txt = dst + ".tmp.txt"
    with open(tmp_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    text_to_pdf(tmp_txt, dst)
    os.remove(tmp_txt)


def csv_to_json(src, dst):
    with open(src, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    with open(dst, "w", encoding="utf-8") as f:
        json.dump(rows, f, indent=2, ensure_ascii=False)


def json_to_csv(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        data = [data]
    if not data:
        open(dst, "w").close()
        return
    keys = list(data[0].keys())
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        writer.writerows(data)


def csv_to_xlsx(src, dst):
    wb = openpyxl.Workbook()
    ws = wb.active
    with open(src, newline="", encoding="utf-8-sig") as f:
        for row in csv.reader(f):
            ws.append(row)
    wb.save(dst)


def xlsx_to_csv(src, dst):
    wb = openpyxl.load_workbook(src)
    ws = wb.active
    with open(dst, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(row)


TARGETS = {
    "image": ["png", "jpg", "bmp", "gif", "webp", "ico", "pdf"],
    "pdf": ["png", "jpg", "txt"],
    "txt": ["pdf"],
    "docx": ["txt", "pdf"],
    "csv": ["json", "xlsx"],
    "json": ["csv"],
    "xlsx": ["csv"],
}


def detect_kind(path):
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    if is_image(path):
        return "image"
    if ext in TARGETS:
        return ext
    return None


def run_conversion(src_paths, target_ext, out_dir):
    results = []
    src_paths = list(src_paths)
    kind = detect_kind(src_paths[0])

    if kind == "image" and target_ext == "pdf" and len(src_paths) > 1:
        dst = os.path.join(out_dir, "merged.pdf")
        images_to_pdf(src_paths, dst)
        results.append(dst)
        return results

    for src in src_paths:
        base = os.path.splitext(os.path.basename(src))[0]
        kind = detect_kind(src)
        dst = os.path.join(out_dir, f"{base}.{target_ext}")

        if kind == "image" and target_ext == "pdf":
            images_to_pdf([src], dst)
        elif kind == "image":
            convert_image(src, dst)
        elif kind == "pdf" and target_ext in ("png", "jpg"):
            results.extend(pdf_to_images(src, out_dir, target_ext))
            continue
        elif kind == "pdf" and target_ext == "txt":
            pdf_to_text(src, dst)
        elif kind == "txt" and target_ext == "pdf":
            text_to_pdf(src, dst)
        elif kind == "docx" and target_ext == "txt":
            docx_to_text(src, dst)
        elif kind == "docx" and target_ext == "pdf":
            docx_to_pdf(src, dst)
        elif kind == "csv" and target_ext == "json":
            csv_to_json(src, dst)
        elif kind == "csv" and target_ext == "xlsx":
            csv_to_xlsx(src, dst)
        elif kind == "json" and target_ext == "csv":
            json_to_csv(src, dst)
        elif kind == "xlsx" and target_ext == "csv":
            xlsx_to_csv(src, dst)
        else:
            raise ValueError(f"Conversion non supportee: {kind} -> {target_ext}")

        results.append(dst)

    return results
